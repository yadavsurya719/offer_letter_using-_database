from docx import Document
from datetime import datetime
import docx
from docx2pdf import convert
#import pandas as pd
import mysql.connector
import pymysql

mydb=mysql.connector.connect(host='localhost',user='root',password='surya7993290146',database='bhagya')
mycursor=mydb.cursor()
print("Database is connected")


def offer_letter(full_name,job_role ,start_date,salary,recipient_name):
    # Create a new Document
    doc = Document()

    # Add a heading
    doc.add_heading('Offer letter', level=1)


    # Add the current date
    current_date = datetime.now().strftime('%Y-%m-%d')
    doc.add_paragraph(f'Date: {current_date}')

    # Add employee details
    doc.add_paragraph(f'Dear {full_name},')
    #doc.add_paragraph(f'Job Role : {job_name}')
    #doc.add_paragraph(f'Salary: {salary}')
    #doc.add_paragraph(f'HR : {hr_name}')
    doc.add_paragraph(f'We at {full_name} are delighted to inform you that you have been selected to our company as a {job_role}. Your joining date is {start_date} and you have to report to the office at 9:AM.Your salary willbe {salary} per month.')
    #doc.add_paragraph(f' {'name'}')
    #doc.add_paragraph(f'Employee Name: {employee_name}')
    #doc.add_paragraph(f'Company: {company_name}')
    #doc.add_paragraph(f'Job Role : {job_name}')
    #doc.add_paragraph(f'Joining date: {start_date}')
    #doc.add_paragraph(f'Salary: {salary}')
    

    # Extract the date part if time information is present
    #start_date = str(start_date).split()[0]
    #end_date = str(end_date).split()[0]


 # Calculate the number of days
    #start_date_dt = datetime.strptime(start_date, '%Y-%m-%d')

    #end_date_dt = datetime.strptime(end_date, '%Y-%m-%d')
    #num_days = (end_date_dt - start_date_dt).days + 1

    #doc.add_paragraph(f'Total Days: {num_days}')

    # Add a closing statement
    #doc.add_paragraph(f'Company : {company_name}')
    
    doc.add_paragraph('Regards,')
    # Add a closing salutation
    #doc.add_paragraph(f'HR : {hr_name}')
    doc.add_paragraph(f'{recipient_name}[HR],Pragmatiq Systems')

    # Save the document
    doc_filename = f'offer_Letter_{full_name.replace(" ","_")}.docx'
    doc.save(doc_filename)

    # Convert the document to PDF
    convert(doc_filename)

   # Connect to MySQL database
mydb = mysql.connector.connect(host='localhost', user='root', password='surya7993290146', database='bhagya')
mycursor = mydb.cursor()



# Select data from the database
mycursor.execute("SELECT * FROM bhagya.data1")
result = mycursor.fetchall()

# Create leave letters for all employees
for row in result:
   offer_letter(row[0], row[1], row[2], row[3], row[4])  # Adjust indices based on your table structure


# Close the database connection
mydb.close()